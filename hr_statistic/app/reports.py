from typing import Any
from docx import Document
from docx.shared import Pt, Inches

import os

class ReportsGenerator:

    def __init__(self, doc) -> None:
        
        self.doc = doc

        self.__output_path = f"{os.getenv('LOC_ROOT_PATH')}" # /opt/render/project/src/hr_statistic/app    # ./hr_statistic/app


    def line_plot_report(self, selected_column1: str, selected_column2: str, intervju_sum: int, start_date: str, end_date: str):
        """### Funkcija koja generise izvestaj za line plot
        :param
        - `selected_column1: str` -> Imerdatoteke
        - `selected_column2: str` -> Ime rukovodilca
        - `intervju_sum: int` -> Broj intervjua 
        - `start_date: str` -> Datum od
        - `end_date: str` -> Datum do
        """
        text = f"""
Rkovodilac {selected_column2} je u tabeli {selected_column1} imao\la ukupan broj intervjua {intervju_sum} za vremenski period od {start_date} do {end_date}.
"""
        section = self.doc.sections[0]

        header = section.header

        h1 = self.doc.add_heading("Izvestaj", 0)
        h1.alignment = 1
        h1_desoration = h1.runs[0].font
        h1_desoration.bold = True
        

        paragraph = self.doc.add_paragraph(text)
        par_decoration = paragraph.runs[0].font
        par_decoration.name = 'Calibri'
        par_decoration.size = Pt(12)  # Postavljamo fonta na 16 Point-a
        

    def pie_plot_report(self, selected_column1: str, start_date: str, end_date: str, data: list):
        """### Funkcija koja generise izvestaj za pie plot
        :param
        - `selected_column1: str` -> Imerdatoteke
        - `start_date: str` -> Datum od
        - `end_date: str` -> Datum do
        - `data: list` -> Lista moze samo da sadrzi update_line_plot i ili update_pie_chart
        """
        # Prelazimo na novi list
        self.doc.add_page_break()

        # Broj kolona i redoa tabele
        row = len(data[0]) + 1
        col = 5

        # Kreiramo zaglavlje dokumenta i dodajemo naslov h1
        section = self.doc.sections[0]
        header = section.header

        h1 = self.doc.add_heading("Izvestaj", 0)
        h1.alignment = 1
        h1_desoration = h1.runs[0].font
        h1_desoration.bold = True
        

        paragraph = self.doc.add_paragraph(f"U tabeli koja sledi u nastavku se prikazuju svi rukovodilci za tabelu {selected_column1} za vremenski period od {start_date} do {end_date}.")
        par_decoration = paragraph.runs[0].font
        par_decoration.name = 'Calibri'
        par_decoration.size = Pt(12)  # Postavljamo fonta na 16 Point-a

        # Dodajemo brake line
        paragraph.add_run().add_break()
        paragraph.add_run().add_break()
        paragraph.add_run().add_break()

        
        table = self.doc.add_table(rows=row, cols=col)
        # Postavljamo prvi reda kao zaglavlje tabele
        table.rows[0].cells[0].paragraphs[0].add_run("Rukovodilac").bold = True
        table.rows[0].cells[1].paragraphs[0].add_run("Ukupan broj intervjua").bold = True
        table.rows[0].cells[2].paragraphs[0].add_run("Ukupan broj primljenih").bold = True
        table.rows[0].cells[3].paragraphs[0].add_run("Pristali na obuku").bold = True
        table.rows[0].cells[4].paragraphs[0].add_run("Odbili obuku").bold = True

        # Popunjavamo celije sa vrednostima, preskacemo prvi red 
        for i in range(0, len(data[0])):  
            # i je kolona j je red
            # Upisujemo vrednost u svaku od celija
            for j in range(len(data)):
                
                cell = table.cell(i+1, j)
                cell.text = f'{data[j][i]}'



    def __call__(self, line_data=None, pie_data=None, both_data=None, report_list=[]) -> Any:
        
        """### Svaki put kada pozovemo instacu cuvamo fajl"""
        # Pre slanja argumenta instanciramo docx
        self.doc = Document()
        # Prolazimo kroz listu i trazimo key
        keys_list = [list(d.keys())[0] for d in report_list]

        # Ako postoji u listi update_line_plot generisi samo njega
        if "update_line_plot" in keys_list and len(keys_list) == 1 and pie_data is None:
            print("update_line_plot")

            line_data = report_list[0][keys_list[0]]

            self.line_plot_report(
                line_data[0], 
                line_data[1], 
                line_data[2], 
                line_data[3], 
                line_data[4]
                )
            
        # Ako postoji u listi update_pie_chart generisi samo njega
        elif "update_pie_chart" in keys_list and len(keys_list) == 1 and line_data is None:
            print("update_pie_chart")

            pie_data = report_list[0][keys_list[0]]

            self.pie_plot_report(
                pie_data[0], 
                pie_data[1], 
                pie_data[2], 
                pie_data[3]
                )
        # Ako postoji u listi update_line_plot i update_pie_chart generisi samo njega
        elif "update_line_plot" in keys_list and "update_pie_chart" in keys_list and len(keys_list) == 2 and line_data is None and pie_data is None:
            print("update_line_plot and update_pie_chart")

            for i in range(0, len(keys_list)):

                if "update_line_plot" in report_list[i]:
                    both_data = report_list[i][keys_list[i]]

                    self.line_plot_report(
                        both_data[0], 
                        both_data[1], 
                        both_data[2], 
                        both_data[3], 
                        both_data[4]
                        )
                    
                elif "update_pie_chart" in report_list[i]:
                    both_data = report_list[i][keys_list[i]]

                    self.pie_plot_report(
                        both_data[0], 
                        both_data[1], 
                        both_data[2], 
                        both_data[3], 
                        )

        new_path = ""
        
        file_name = "izvestaj.docx"
            
        path_used = new_path.endswith("/izvestaj.docx")

        if path_used == True:

            splited_str = self.__output_path.split("/")

            splited_str.pop(-1)

            new_path = "/".join(splited_str)

            new_path +=  "/" + file_name

            self.doc.save(new_path)

            return new_path

        else:

            new_path += self.__output_path + "/" + file_name

            self.doc.save(new_path)

            return new_path
        
        



# dic = [
#     {'update_line_plot': ['Komercijala', 'Nemanja Božinovski', 5, '01-06-2023', '30-06-2023']}, 
#     {'update_pie_chart': ['Komercijala', '01-06-2023', '30-06-2023', [
#         ['Voja', 'Nemanja Božinovski', 'Goran', 'Peđa', 'Đole', 'Čeda', 'Aleksa'], 
#         [6, 5, 12, 4, 1, 2, 2], 
#         [6, 0, 7, 2, 1, 0, 0], 
#         [5, 0, 3, 2, 1, 0, 0], 
#         [1, 0, 4, 0, 0, 0, 0]]
#         ]
#     }
# ]
# keys_list = [list(d.keys())[0] for d in dic]
# print(keys_list)
# for i in range(0, len(dic)):
#     if "update_line_plot" in dic[i]:
#         dic[i].pop("update_line_plot", None)
#         # print()


 
# print(len(dic))
# d = ReportsGenerator(Document())


# d(report_list=dic)
# d(report_list=dic)
# d(report_list=dic)


# d.pie_plot_report("Komercijala", "01-09-2023", "31-09-2023", 0)



